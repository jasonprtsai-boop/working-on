from __future__ import annotations

import json
import sys
import zipfile
from pathlib import Path

from docx import Document


DOCX_PATHS = [
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\分段\第一段.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\分段\第二章.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\分段\第三章.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\分段\第四章 系統實作與初步測試規劃.docx",
    r"C:\Users\user\Desktop\01_file_inventory.md (line 1).docx",
    r"C:\Users\user\Desktop\02_data_flow.md (line 1).docx",
    r"C:\Users\user\Desktop\03_architecture_and_runtime.md (line 1)：.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\新的流程圖\chap plus.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\新的流程圖\codex 建議.docx",
    r"C:\Users\user\Desktop\專題\ppt & 書面報告\新的流程圖\圖 建議.docx",
]

ZIP_PATH = Path(r"C:\Users\user\Desktop\專題\專題程式\壓縮\code-19.zip")


def paragraph_text(doc: Document) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue
        rows.append({"style": para.style.name if para.style else "", "text": text})
    return rows


def table_text(doc: Document) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([" ".join(cell.text.split()) for cell in row.cells])
        if rows:
            tables.append(rows)
    return tables


def extract_docx(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = paragraph_text(doc)
    tables = table_text(doc)
    return {
        "path": str(path),
        "name": path.name,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "image_count": len(doc.inline_shapes),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def zip_listing(zip_path: Path) -> dict:
    with zipfile.ZipFile(zip_path) as zf:
        infos = zf.infolist()
        names = [info.filename for info in infos]
        top_dirs: dict[str, int] = {}
        extensions: dict[str, int] = {}
        for name in names:
            parts = name.replace("\\", "/").split("/")
            top = parts[0] if len(parts) > 1 else "."
            top_dirs[top] = top_dirs.get(top, 0) + 1
            ext = Path(name).suffix.lower() or "[none]"
            extensions[ext] = extensions.get(ext, 0) + 1
        return {
            "path": str(zip_path),
            "file_count": len(infos),
            "top_dirs": dict(sorted(top_dirs.items(), key=lambda item: (-item[1], item[0]))),
            "extensions": dict(sorted(extensions.items(), key=lambda item: (-item[1], item[0]))),
            "sample_files": names[:300],
        }


def write_markdown(output_dir: Path, docs: list[dict], zip_info: dict) -> None:
    lines: list[str] = ["# Report Source Extract", ""]
    for item in docs:
        lines.extend(
            [
                f"## {item['name']}",
                f"- Path: `{item['path']}`",
                f"- Paragraphs: {item['paragraph_count']}",
                f"- Tables: {item['table_count']}",
                f"- Images: {item['image_count']}",
                "",
            ]
        )
        for para in item["paragraphs"]:
            style = para["style"]
            text = para["text"]
            if style.startswith("Heading"):
                lines.append(f"### {text}")
            else:
                lines.append(text)
            lines.append("")
        for table_index, table in enumerate(item["tables"], start=1):
            lines.append(f"Table {table_index}:")
            for row in table:
                lines.append(" | ".join(row))
            lines.append("")

    lines.extend(
        [
            "## code-19.zip inventory",
            f"- Path: `{zip_info['path']}`",
            f"- Files: {zip_info['file_count']}",
            "",
            "### Top directories",
        ]
    )
    for name, count in zip_info["top_dirs"].items():
        lines.append(f"- {name}: {count}")
    lines.append("")
    lines.append("### Extensions")
    for ext, count in zip_info["extensions"].items():
        lines.append(f"- {ext}: {count}")
    lines.append("")
    lines.append("### Sample files")
    for name in zip_info["sample_files"]:
        lines.append(f"- {name}")

    (output_dir / "source_extract.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> int:
    output_dir = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("report_diagrams/source_extract")
    output_dir.mkdir(parents=True, exist_ok=True)
    docs = [extract_docx(Path(path)) for path in DOCX_PATHS]
    zip_info = zip_listing(ZIP_PATH)
    (output_dir / "source_extract.json").write_text(
        json.dumps({"documents": docs, "zip": zip_info}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    write_markdown(output_dir, docs, zip_info)
    print(json.dumps({"output_dir": str(output_dir), "documents": len(docs), "zip_files": zip_info["file_count"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
